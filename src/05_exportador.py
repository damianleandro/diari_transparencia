import docx
from docx.shared import Inches
import os

class ExportadorWord:
    def __init__(self):
        self.doc = docx.Document()
        
    def generar_document(self, text_noticia, text_podcast, ruta_imatge, nom_arxiu="Reportatge_Multimèdia_Sequera.docx"):
        print("📝 Maquetant el document Word amb Notícia i Guió de Podcast...")
        
        # --- PÀGINA 1: EL REPORTATGE ESCRIT ---
        self.doc.add_heading("Agència de Notícies IA - El Cronista de Dades", 0)
        
        if os.path.exists(ruta_imatge):
            self.doc.add_picture(ruta_imatge, width=Inches(6.0))
            paragraf_imatge = self.doc.paragraphs[-1]
            paragraf_imatge.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph() 
        
        for paragraf in text_noticia.split('\n'):
            if paragraf.strip(): 
                if paragraf.startswith('**') and paragraf.endswith('**'):
                    self.doc.add_heading(paragraf.replace('**', ''), level=1)
                else:
                    self.doc.add_paragraph(paragraf.strip())
                    
        # --- PÀGINA 2: EL GUIÓ DE RÀDIO / PODCAST ---
        self.doc.add_page_break() # Salt de pàgina
        self.doc.add_heading("🎙️ Guió de Ràdio: La Dada Clara", level=1)
        self.doc.add_paragraph("Guió autogenerat preparat per a locució humana o Text-To-Speech.\n")
        
        for linia in text_podcast.split('\n'):
            if linia.strip():
                p = self.doc.add_paragraph()
                # Si la línia comença pel nom del locutor, el posem en negreta
                if linia.startswith("MARC:") or linia.startswith("ANNA:"):
                    parts = linia.split(":", 1)
                    p.add_run(parts[0] + ":").bold = True
                    if len(parts) > 1:
                        p.add_run(parts[1])
                else:
                    p.add_run(linia.strip())

        # Peu de pàgina final
        peu = self.doc.add_paragraph("\nDocument i guió generats automàticament a partir de Dades Obertes de la Generalitat de Catalunya.")
        peu.style = 'Intense Quote'
        
        # Guardar l'arxiu
        ruta_arrel = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        ruta_final = os.path.join(ruta_arrel, nom_arxiu)
        self.doc.save(ruta_final)
        
        return ruta_final